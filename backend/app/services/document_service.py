"""Phase 3 — Document Ingestion Pipeline (text extraction / ETL).

Extracts raw text from uploaded documents (PDF, DOCX, TXT, HTML), updates the
document processing status, and — once extraction succeeds — automatically
triggers the Phase 4 chunking pipeline.
"""
import logging

import fitz  # PyMuPDF
import docx  # python-docx
from bs4 import BeautifulSoup
from sqlalchemy.orm import Session

from app.models.document import Document, DocumentStatus, DocumentType
from app.services.text_processing_service import process_chunks

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str):
    """Extract text from a PDF using PyMuPDF (fitz).

    Returns a tuple of ``(text, page_count)``.
    """
    text_parts = []
    with fitz.open(file_path) as pdf:
        page_count = pdf.page_count
        for page in pdf:
            text_parts.append(page.get_text())
    return "\n".join(text_parts), page_count


def extract_text_from_docx(file_path: str):
    """Extract text from a DOCX file using python-docx.

    Returns ``(text, page_count)`` — page_count is ``None`` (not available
    from the DOCX format without rendering).
    """
    document = docx.Document(file_path)
    paragraphs = [p.text for p in document.paragraphs]
    # Include table cell text so tabular content is not lost.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs), None


def extract_text_from_txt(file_path: str):
    """Read a plain-text file. Returns ``(text, None)``."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read(), None


def extract_text_from_html(file_path: str):
    """Extract visible text from an HTML file using BeautifulSoup.

    Returns ``(text, None)``.
    """
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    # Drop script/style content before extracting text.
    for tag in soup(["script", "style"]):
        tag.decompose()
    return soup.get_text(separator="\n"), None


_EXTRACTORS = {
    DocumentType.PDF: extract_text_from_pdf,
    DocumentType.DOCX: extract_text_from_docx,
    DocumentType.TXT: extract_text_from_txt,
    DocumentType.HTML: extract_text_from_html,
}


def process_document(doc_id, db: Session) -> Document:
    """Orchestrate extraction for a single document.

    Marks the document PROCESSING, runs the type-specific extractor, stores the
    raw text + page count, then triggers chunking and marks it PROCESSED. On any
    failure the document is marked FAILED with the error message recorded.
    """
    document = db.query(Document).filter(Document.id == doc_id).first()
    if document is None:
        raise ValueError(f"Document {doc_id} not found")

    document.status = DocumentStatus.PROCESSING
    document.error_message = None
    db.commit()

    try:
        extractor = _EXTRACTORS.get(document.document_type)
        if extractor is None:
            raise ValueError(
                f"Unsupported document type: {document.document_type}"
            )

        text, page_count = extractor(document.file_path)

        document.raw_text = text
        document.page_count = page_count
        db.commit()

        # Phase 4: automatically chunk the freshly extracted text.
        process_chunks(document.id, db)

        document.status = DocumentStatus.PROCESSED
        document.error_message = None
        db.commit()
        db.refresh(document)
        logger.info("Document %s processed successfully", doc_id)
    except Exception as exc:  # noqa: BLE001 — record any extraction failure
        db.rollback()
        document = db.query(Document).filter(Document.id == doc_id).first()
        if document is not None:
            document.status = DocumentStatus.FAILED
            document.error_message = str(exc)[:1000]
            db.commit()
        logger.exception("Failed to process document %s: %s", doc_id, exc)

    return document
